import os
import streamlit as st
from docx import Document
from datetime import datetime
from docx.oxml.ns import qn
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import uuid
import tempfile
import platform
import subprocess
import shutil

# Configuration for file paths
pdf_dir = "/tmp"
os.makedirs(pdf_dir, exist_ok=True)

PROPOSAL_CONFIG = {
    "Manychats + CRM Automation - 550 USD": {
        "template": "HVT Proposal - AI Automations.docx",
        "special_fields": [("VDate", "<<")],
        "team_type": "hvt_ai"
    },
    "Manychats + CRM Automation - Custom Price": {
        "template": "HVT Proposal - AI Automations - Custom Price.docx",
        "special_fields": [("VDate", "<<")],
        "team_type": "hvt_ai_custom_price"
    },
    "Internship Offer Letter": {
        "template": "Offer Letter.docx",
        "special_fields": [],
        "team_type": "offer_letter"
    }
}

def apply_run_formatting(new_run, source_run):
    if source_run is None:
        return
    if source_run.font.name:
        new_run.font.name = source_run.font.name
        rPr = new_run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), source_run.font.name)
    if source_run.font.size:
        new_run.font.size = source_run.font.size
    if source_run.font.color and source_run.font.color.rgb:
        new_run.font.color.rgb = source_run.font.color.rgb
    new_run.bold = source_run.bold
    new_run.italic = source_run.italic
    new_run.underline = source_run.underline

def replace_placeholder(paragraph, placeholder, value):
    if placeholder not in paragraph.text:
        return False
    if not paragraph.runs:
        paragraph.text = paragraph.text.replace(placeholder, str(value))
        return True

    runs = paragraph.runs
    full_text = ''.join([run.text for run in runs])
    
    if placeholder not in full_text:
        return False

    start_idx = full_text.find(placeholder)
    end_idx = start_idx + len(placeholder)
    
    for run in runs:
        run.text = ""
    
    before = full_text[:start_idx]
    after = full_text[end_idx:]
    
    if before:
        new_run = paragraph.add_run(before)
        apply_run_formatting(new_run, runs[0])
    
    new_run = paragraph.add_run(str(value))
    apply_run_formatting(new_run, runs[0])
    
    if after:
        new_run = paragraph.add_run(after)
        apply_run_formatting(new_run, runs[-1])
    
    return True

def process_document(doc, placeholders):
    for paragraph in doc.paragraphs:
        if not paragraph.text:
            continue
        for ph, value in placeholders.items():
            replace_placeholder(paragraph, ph, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.tables:
                    for nested_table in cell.tables:
                        for nested_row in nested_table.rows:
                            for nested_cell in nested_row.cells:
                                for para in nested_cell.paragraphs:
                                    if not para.text:
                                        continue
                                    for ph, value in placeholders.items():
                                        replace_placeholder(para, ph, value)
                for para in cell.paragraphs:
                    if not para.text:
                        continue
                    for ph, value in placeholders.items():
                        replace_placeholder(para, ph, value)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return doc

def get_hvt_ai_team_details():
    st.subheader("Team Composition")
    team_roles = {
        "Project Manager": "P1",
        "Frontend Developers": "F1",
        "UI/UX Members": "U1",
        "AI/ML Developers": "A1",
        "Business Analyst": "B1",
        "AWS Developer": "AD1",
        "Backend Developers": "BD1",
        "System Architect": "S1"
    }
    
    team_details = {}
    cols = st.columns(2)
    
    for idx, (role, placeholder) in enumerate(team_roles.items()):
        with cols[idx % 2]:
            team_details[f"<<{placeholder}>>"] = st.number_input(
                f"{role} Count:",
                min_value=0,
                step=1,
                key=f"hvt_team_{placeholder}"
            )
    return team_details

def validate_phone_number(country, number):
    if not number:
        return True
    if country.lower() == "india":
        return number.startswith("+91")
    return number.startswith("+1")


def convert_to_pdf(doc_path, pdf_path):
    """Converts a Word document to a flattened PDF."""
    doc_path = os.path.abspath(doc_path)
    pdf_path = os.path.abspath(pdf_path)

    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"Word document not found: {doc_path}")

    with tempfile.TemporaryDirectory() as temp_dir:
        # LibreOffice automatically generates a PDF with the same name as the input file
        expected_pdf_path = os.path.join(temp_dir, os.path.basename(doc_path).replace('.docx', '.pdf'))

        if platform.system() == "Windows":
            try:
                import comtypes.client
                import pythoncom
                pythoncom.CoInitialize()
                word = comtypes.client.CreateObject("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(doc_path)
                doc.SaveAs(expected_pdf_path, FileFormat=17)
                doc.Close()
                word.Quit()
                pythoncom.CoUninitialize()

                if not os.path.exists(expected_pdf_path):
                    raise FileNotFoundError("PDF conversion failed.")

                shutil.move(expected_pdf_path, pdf_path)

            except Exception as e:
                raise Exception(f"Error in Windows Word to PDF conversion: {e}")

        else:
            try:
                subprocess.run(
                    ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', temp_dir, doc_path],
                    check=True
                )

                # Find the actual PDF file
                generated_files = os.listdir(temp_dir)
                pdf_files = [f for f in generated_files if f.endswith('.pdf')]

                if not pdf_files:
                    raise FileNotFoundError("LibreOffice PDF conversion failed: No PDF file generated.")

                generated_pdf_path = os.path.join(temp_dir, pdf_files[0])  # Pick the first PDF file

                shutil.move(generated_pdf_path, pdf_path)

            except subprocess.CalledProcessError as e:
                raise Exception(f"Error in LibreOffice conversion: {e}")

def generate_document():
    st.title("Document Generator")

    
    # Get user selection **before** accessing config
    selected_proposal = st.selectbox("Select Document Type", list(PROPOSAL_CONFIG.keys()))
    config = PROPOSAL_CONFIG[selected_proposal]  # Now, config is correctly assigned

    # Now safely access config["template"]
    template_path = config["template"]

    if not os.path.exists(template_path):
        st.error(f"Template file not found: {template_path}")
        raise FileNotFoundError(f"Template file not found: {template_path}")

    if 'generated_files' not in st.session_state:
        st.session_state.generated_files = {}

    placeholders = {}
    if selected_proposal == "Internship Offer Letter":
        st.subheader("Candidate Information")
        placeholders.update({
            "<<E-Name>>": st.text_input("Candidate Name:"),
            "<<Job>>": st.selectbox("Job Role", ["UI UX", "AI Automations", "Software Developer", "Sales"]),
            "<<S-Date>>": st.date_input("Start Date").strftime("%d %B, %Y"),
            "<<Stipend>>": f"{st.number_input('Stipend (₹)', min_value=0):,}",
            "<<Months>>": st.number_input("Duration (Months)", min_value=1),
            "<<Date>>": datetime.today().strftime("%d %B, %Y")
        })
    else:
        st.subheader("Client Details")
        col1, col2 = st.columns(2)
        with col1:
            client_name = st.text_input("Client Name:")
            client_email = st.text_input("Email:")
        with col2:
            country = st.text_input("Country:")
            client_number = st.text_input("Phone Number:")
        
        st.subheader("Date Information")
        date_col1, date_col2 = st.columns(2)
        with date_col1:
            date_field = st.date_input("Proposal Date", datetime.today())
        with date_col2:
            validation_date = st.date_input("Validation Date", datetime.today())

        placeholders.update({
            "<<Client Name>>": client_name,
            "<<Client Email>>": client_email,
            "<<Client Number>>": client_number,
            "<<Country>>": country,
            "<<Date>>": date_field.strftime("%d %B, %Y"),
            "<<D-Date>>": date_field.strftime("%d %B, %Y"),
            "<<VDate>>": validation_date.strftime("%d-%m-%Y")
        })

        if "hvt_ai" in config["team_type"]:
            placeholders.update(get_hvt_ai_team_details())

        if "custom_price" in config["team_type"]:
            st.subheader("Pricing Details")
            pricing = {
                "<<P01>>": st.number_input("Manychats Setup (USD)", min_value=0),
                "<<P02>>": st.number_input("Make Automations (USD)", min_value=0),
                "<<A-Price>>": st.number_input("Annual Maintenance (USD)", min_value=0)
            }
            placeholders.update(pricing)
            placeholders["<<T-Price>>"] = f"{sum(pricing.values()):,}"

    if st.button("Generate Documents"):
        if selected_proposal != "Internship Offer Letter":
            if not validate_phone_number(placeholders["<<Country>>"], placeholders["<<Client Number>>"]):
                st.error("Invalid phone number format for selected country")
                return

        unique_id = uuid.uuid4().hex[:8]
        base_name = f"{selected_proposal.replace(' ', '_')}_{unique_id}"
        doc_filename = f"{base_name}.docx"
        pdf_filename = f"{base_name}.pdf"

        try:

            # Process Word document
            doc = Document(template_path)
            doc = process_document(doc, placeholders)
            doc_path = f"{base_name}.docx"
            doc.save(doc_path)

            # Convert to PDF with retries
            pdf_path = doc_path.replace(".docx", ".pdf")
            convert_to_pdf(doc_path, pdf_path)

            # Verify PDF file exists before storing
            if not os.path.exists(pdf_path):
                raise FileNotFoundError(f"PDF file not found at: {pdf_path}")

            # Store in session state
            with open(doc_path, "rb") as f:
                st.session_state.generated_files['doc'] = f.read()
            with open(pdf_path, "rb") as f:
                st.session_state.generated_files['pdf'] = f.read()

            st.session_state.generated_files['doc_name'] = doc_filename
            st.session_state.generated_files['pdf_name'] = pdf_filename
            st.success("Documents generated successfully!")

        except Exception as e:
            st.error(f"Generation failed: {str(e)}")

    if 'doc' in st.session_state.generated_files:
        st.markdown("---")
        st.subheader("Download Documents")
        
        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                label="Download Word Document",
                data=st.session_state.generated_files['doc'],
                file_name=st.session_state.generated_files['doc_name'],
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        with col2:
            st.download_button(
                label="Download PDF Document",
                data=st.session_state.generated_files['pdf'],
                file_name=st.session_state.generated_files['pdf_name'],
                mime="application/pdf"
            )

def main():
    generate_document()

if __name__ == "__main__":
    main()
